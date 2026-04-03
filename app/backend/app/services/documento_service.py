"""
Serviço de Geração de Documentos - DOCX e PDF
"""

import io
import logging
from datetime import datetime
from typing import Optional, Dict, Any
from pathlib import Path

from app.core.llm import get_llm_engine as get_llm_client

logger = logging.getLogger(__name__)


class DocumentoService:
    """
    Serviço para geração de documentos jurídicos em DOCX e PDF
    """
    
    def __init__(self):
        self.templates_path = Path("templates/documentos")
        self.llm_client = None

    async def _ensure_llm(self):
        if not self.llm_client:
            self.llm_client = await get_llm_client()

    async def analisar_documento(self, conteudo: str, tipo_analise: str = "geral") -> Dict[str, Any]:
        """
        Analisa documento jurídico usando LLM
        """
        await self._ensure_llm()
        
        prompts = {
            "geral": "Analise o documento jurídico abaixo e forneça um resumo, pontos principais e possíveis riscos.",
            "riscos": "Analise o documento jurídico abaixo focando EXCLUSIVAMENTE em riscos legais e contratuais.",
            "melhorias": "Sugira melhorias na redação e fundamentação do documento jurídico abaixo."
        }
        
        prompt_base = prompts.get(tipo_analise, prompts["geral"])
        prompt = f"{prompt_base}\n\n=== DOCUMENTO ===\n{conteudo[:10000]}" # Limite de chars
        
        try:
            response = await self.llm_client.processar_consulta_juridica(
                consulta=prompt,
                contexto_rag=None # Análise pura do documento
            )
            
            return {
                "analise": response.content,
                "tokens": response.tokens_used,
                "modelo": response.model,
                "timestamp": datetime.now().isoformat()
            }
        except Exception as e:
            logger.error(f"Erro na análise de documento: {e}")
            return {"error": str(e)}

    async def gerar_docx(
        self,
        conteudo: str,
        titulo: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Gera documento DOCX a partir do conteúdo
        
        Args:
            conteudo: Texto da petição/documento
            titulo: Título do documento
            metadata: Metadados adicionais
            
        Returns:
            Bytes do arquivo DOCX
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            logger.warning("python-docx não instalado, usando geração simplificada")
            return self._gerar_docx_simples(conteudo, titulo)
        
        doc = Document()
        
        # Configurar margens
        for section in doc.sections:
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)
        
        # Título
        titulo_para = doc.add_paragraph()
        titulo_run = titulo_para.add_run(titulo.upper())
        titulo_run.bold = True
        titulo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espaço
        
        # Metadados se fornecidos
        if metadata:
            if metadata.get("numero_processo"):
                meta_para = doc.add_paragraph()
                meta_run = meta_para.add_run(f"Processo nº: {metadata['numero_processo']}")
                meta_run.font.size = Pt(12)
                meta_run.font.name = 'Times New Roman'
            
            if metadata.get("autor"):
                meta_para = doc.add_paragraph()
                meta_run = meta_para.add_run(f"Autor: {metadata['autor']}")
                meta_run.font.size = Pt(12)
                meta_run.font.name = 'Times New Roman'
            
            if metadata.get("reu"):
                meta_para = doc.add_paragraph()
                meta_run = meta_para.add_run(f"Réu: {metadata['reu']}")
                meta_run.font.size = Pt(12)
                meta_run.font.name = 'Times New Roman'
            
            doc.add_paragraph()  # Espaço
        
        # Conteúdo principal
        paragrafos = conteudo.split('\n\n')
        
        for paragrafo in paragrafos:
            if not paragrafo.strip():
                continue
            
            # Verificar se é título de seção (linha que termina com :)
            linhas = paragrafo.strip().split('\n')
            
            for linha in linhas:
                if not linha.strip():
                    continue
                
                p = doc.add_paragraph()
                
                # Detectar títulos de seção
                if linha.strip().endswith(':') and len(linha) < 50:
                    run = p.add_run(linha.strip())
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Detectar numeração romana (I., II., III., etc.)
                elif linha.strip()[:3] in ['I. ', 'II.', 'III', 'IV.', 'V. ', 'VI.', 'VII', 'VII', 'IX.', 'X. ']:
                    run = p.add_run(linha.strip())
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
                
                else:
                    run = p.add_run(linha.strip())
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    # Recuo primeira linha
                    p.paragraph_format.first_line_indent = Cm(1.25)
        
        # Data e local
        doc.add_paragraph()
        data_para = doc.add_paragraph()
        data_atual = datetime.now().strftime("%d de %B de %Y")
        # Traduzir meses
        meses = {
            'January': 'janeiro', 'February': 'fevereiro', 'March': 'março',
            'April': 'abril', 'May': 'maio', 'June': 'junho',
            'July': 'julho', 'August': 'agosto', 'September': 'setembro',
            'October': 'outubro', 'November': 'novembro', 'December': 'dezembro'
        }
        for en, pt in meses.items():
            data_atual = data_atual.replace(en, pt)
        
        cidade = metadata.get("cidade", "São Paulo") if metadata else "São Paulo"
        data_run = data_para.add_run(f"{cidade}, {data_atual}.")
        data_run.font.size = Pt(12)
        data_run.font.name = 'Times New Roman'
        data_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Assinatura
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        assinatura_para = doc.add_paragraph()
        assinatura_para.add_run("_" * 40)
        assinatura_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        advogado_para = doc.add_paragraph()
        advogado = metadata.get("advogado", "Advogado(a)") if metadata else "Advogado(a)"


        oab = metadata.get("oab", "OAB/XX 000.000") if metadata else "OAB/XX 000.000"
        # Garantir que o texto do advogado e OAB sejam adicionados corretamente
        advogado_run = advogado_para.add_run(f"{advogado}\n{oab}")
        advogado_run.font.size = Pt(12)
        advogado_run.font.name = 'Times New Roman'
        advogado_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Salvar em bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.read()
    
    def _gerar_docx_simples(self, conteudo: str, titulo: str) -> bytes:
        """Geração DOCX simplificada quando python-docx não está disponível"""
        # Retorna texto simples como fallback
        texto = f"{titulo.upper()}\n\n{conteudo}"
        return texto.encode('utf-8')
    
    async def gerar_pdf(
        self,
        conteudo: str,
        titulo: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Gera documento PDF a partir do conteúdo
        
        Args:
            conteudo: Texto da petição/documento
            titulo: Título do documento
            metadata: Metadados adicionais
            
        Returns:
            Bytes do arquivo PDF
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_RIGHT
        except ImportError:
            logger.warning("reportlab não instalado, usando conversão via DOCX")
            return await self._converter_docx_para_pdf(conteudo, titulo, metadata)
        
        buffer = io.BytesIO()
        
        # Criar documento
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=3*cm,
            topMargin=3*cm,
            bottomMargin=2*cm
        )
        
        # Estilos
        styles = getSampleStyleSheet()
        
        titulo_style = ParagraphStyle(
            'TituloPeticao',
            parent=styles['Heading1'],
            fontSize=14,
            fontName='Times-Bold',
            alignment=TA_CENTER,
            spaceAfter=20
        )
        
        normal_style = ParagraphStyle(
            'NormalJuridico',
            parent=styles['Normal'],
            fontSize=12,
            fontName='Times-Roman',
            alignment=TA_JUSTIFY,
            firstLineIndent=1.25*cm,
            spaceBefore=6,
            spaceAfter=6
        )
        
        secao_style = ParagraphStyle(
            'SecaoJuridica',
            parent=styles['Normal'],
            fontSize=12,
            fontName='Times-Bold',
            alignment=TA_JUSTIFY,
            spaceBefore=12,
            spaceAfter=6
        )
        
        direita_style = ParagraphStyle(
            'DireitaJuridico',
            parent=styles['Normal'],
            fontSize=12,
            fontName='Times-Roman',
            alignment=TA_RIGHT
        )
        
        centro_style = ParagraphStyle(
            'CentroJuridico',
            parent=styles['Normal'],
            fontSize=12,
            fontName='Times-Roman',
            alignment=TA_CENTER
        )
        
        # Elementos do documento
        elements = []
        
        # Título
        elements.append(Paragraph(titulo.upper(), titulo_style))
        elements.append(Spacer(1, 20))
        
        # Metadados
        if metadata:
            if metadata.get("numero_processo"):
                elements.append(Paragraph(
                    f"<b>Processo nº:</b> {metadata['numero_processo']}", 
                    normal_style
                ))
            if metadata.get("autor"):
                elements.append(Paragraph(
                    f"<b>Autor:</b> {metadata['autor']}", 
                    normal_style
                ))
            if metadata.get("reu"):
                elements.append(Paragraph(
                    f"<b>Réu:</b> {metadata['reu']}", 
                    normal_style
                ))
            elements.append(Spacer(1, 20))
        
        # Conteúdo
        paragrafos = conteudo.split('\n\n')
        
        for paragrafo in paragrafos:
            if not paragrafo.strip():
                continue
            
            linhas = paragrafo.strip().split('\n')
            
            for linha in linhas:
                if not linha.strip():
                    continue
                
                # Escapar caracteres especiais do ReportLab
                linha_escaped = linha.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # Detectar títulos de seção
                if linha.strip().endswith(':') and len(linha) < 50:
                    elements.append(Paragraph(f"<b>{linha_escaped}</b>", secao_style))
                elif linha.strip()[:3] in ['I. ', 'II.', 'III', 'IV.', 'V. ', 'VI.', 'VII', 'VII', 'IX.', 'X. ']:
                    elements.append(Paragraph(f"<b>{linha_escaped}</b>", secao_style))
                else:
                    elements.append(Paragraph(linha_escaped, normal_style))
        
        # Data e local
        elements.append(Spacer(1, 30))
        data_atual = datetime.now().strftime("%d de %B de %Y")
        meses = {
            'January': 'janeiro', 'February': 'fevereiro', 'March': 'março',
            'April': 'abril', 'May': 'maio', 'June': 'junho',
            'July': 'julho', 'August': 'agosto', 'September': 'setembro',
            'October': 'outubro', 'November': 'novembro', 'December': 'dezembro'
        }
        for en, pt in meses.items():
            data_atual = data_atual.replace(en, pt)
        
        cidade = metadata.get("cidade", "São Paulo") if metadata else "São Paulo"
        elements.append(Paragraph(f"{cidade}, {data_atual}.", direita_style))
        
        # Assinatura
        elements.append(Spacer(1, 50))
        elements.append(Paragraph("_" * 50, centro_style))
        
        advogado = metadata.get("advogado", "Advogado(a)") if metadata else "Advogado(a)"
        oab = metadata.get("oab", "OAB/XX 000.000") if metadata else "OAB/XX 000.000"
        elements.append(Paragraph(f"{advogado}<br/>{oab}", centro_style))
        
        # Gerar PDF
        doc.build(elements)
        buffer.seek(0)
        
        return buffer.read()
    
    async def _converter_docx_para_pdf(
        self,
        conteudo: str,
        titulo: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """Fallback: gera texto simples quando reportlab não está disponível"""
        texto = f"{titulo.upper()}\n\n{conteudo}"
        return texto.encode('utf-8')
    
    async def extrair_texto_pdf(self, pdf_bytes: bytes) -> str:
        """
        Extrai texto de um arquivo PDF
        
        Args:
            pdf_bytes: Bytes do arquivo PDF
            
        Returns:
            Texto extraído
        """
        try:
            import pypdf
            
            reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
            texto = ""
            
            for page in reader.pages:
                texto += page.extract_text() + "\n"
            
            return texto.strip()
            
        except ImportError:
            logger.warning("pypdf não instalado")
            return ""
        except Exception as e:
            logger.error(f"Erro ao extrair texto do PDF: {e}")
            return ""
    
    async def extrair_texto_docx(self, docx_bytes: bytes) -> str:
        """
        Extrai texto de um arquivo DOCX
        
        Args:
            docx_bytes: Bytes do arquivo DOCX
            
        Returns:
            Texto extraído
        """
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(docx_bytes))
            texto = ""
            
            for para in doc.paragraphs:
                texto += para.text + "\n"
            
            return texto.strip()
            
        except ImportError:
            logger.warning("python-docx não instalado")
            return ""
        except Exception as e:
            logger.error(f"Erro ao extrair texto do DOCX: {e}")
            return ""


# Instância global
_documento_service = None


def get_documento_service() -> DocumentoService:
    """Dependency para obter DocumentoService"""
    global _documento_service
    if _documento_service is None:
        _documento_service = DocumentoService()
    return _documento_service